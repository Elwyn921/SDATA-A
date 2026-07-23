import json
from datetime import date

from docx import Document

from satellite_news.reporting.periodic_documents import (
    build_previous_day_payload,
    build_weekly_payload,
    generate_previous_day_document,
    generate_weekly_document,
)


def sample_catalog() -> dict:
    return {
        "generated_at": "2026-07-23T03:00:00Z",
        "items": [
            {
                "id": "today",
                "company_id": "spacex",
                "company_name": "SpaceX",
                "title": "Today item",
                "published_at": "2026-07-23T01:00:00Z",
                "source": {"source_name": "Source A"},
                "quality": {"event_type": "launch", "company_relevance_score": 0.9},
            },
            {
                "id": "yesterday",
                "company_id": "landspace",
                "company_name": "蓝箭航天",
                "title": "蓝箭航天完成试验",
                "published_at": "2026-07-22T02:00:00Z",
                "source": {"source_name": "Source B"},
                "quality": {"event_type": "product", "company_relevance_score": 0.8},
            },
            {
                "id": "week",
                "company_id": "landspace",
                "company_name": "蓝箭航天",
                "title": "蓝箭航天此前完成发动机试验",
                "published_at": "2026-07-19T02:00:00Z",
                "source": {"source_name": "Source C"},
                "quality": {"event_type": "product", "company_relevance_score": 0.8},
            },
        ],
    }


def sample_timeline() -> dict:
    return {
        "events": [
            {
                "event_id": "current",
                "company_id": "landspace",
                "company_name": "蓝箭航天",
                "event_type": "product",
                "headline": "蓝箭航天完成发动机试验",
                "latest_at": "2026-07-22T02:00:00Z",
                "importance_score": 80,
                "source_count": 1,
                "source_names": ["Source B"],
                "articles": [{"published_at": "2026-07-22T02:00:00Z"}],
            },
            {
                "event_id": "older",
                "company_id": "landspace",
                "company_name": "蓝箭航天",
                "event_type": "product",
                "headline": "蓝箭航天开展发动机试验",
                "latest_at": "2026-07-01T02:00:00Z",
                "importance_score": 60,
                "source_count": 1,
                "articles": [{"published_at": "2026-07-01T02:00:00Z"}],
            },
        ]
    }


def test_previous_day_payload_uses_exact_previous_natural_day():
    payload = build_previous_day_payload(sample_catalog(), as_of=date(2026, 7, 23))

    assert payload["report_date"] == "2026-07-22"
    assert payload["total_items"] == 1
    assert payload["top_items"][0]["item_id"] == "yesterday"


def test_weekly_payload_covers_seven_days_and_adds_related_history():
    payload = build_weekly_payload(
        sample_catalog(),
        sample_timeline(),
        as_of=date(2026, 7, 23),
    )

    assert payload["period_start"] == "2026-07-17"
    assert payload["period_end"] == "2026-07-23"
    assert payload["total_items"] == 3
    assert payload["event_count"] == 1
    assert payload["related_history"][0]["historical_event_id"] == "older"


def test_generators_write_editable_docx_and_json(tmp_path):
    catalog_path = tmp_path / "catalog.json"
    timeline_path = tmp_path / "timeline.json"
    catalog_path.write_text(json.dumps(sample_catalog(), ensure_ascii=False), encoding="utf-8")
    timeline_path.write_text(json.dumps(sample_timeline(), ensure_ascii=False), encoding="utf-8")

    daily_payload, daily_outputs = generate_previous_day_document(
        catalog_path=catalog_path,
        local_root=tmp_path / "daily",
        publish_root=tmp_path / "published-daily",
        as_of=date(2026, 7, 23),
    )
    weekly_payload, weekly_outputs = generate_weekly_document(
        catalog_path=catalog_path,
        event_timeline_path=timeline_path,
        local_root=tmp_path / "weekly",
        publish_root=tmp_path / "published-weekly",
        as_of=date(2026, 7, 23),
    )

    assert daily_payload["total_items"] == 1
    assert weekly_payload["total_items"] == 3
    assert daily_outputs.latest_docx.exists()
    assert weekly_outputs.published_docx.exists()
    assert daily_outputs.archived_json.exists()
    assert weekly_outputs.archived_json.exists()
    assert "商业航天新闻日报" in "\n".join(
        paragraph.text for paragraph in Document(daily_outputs.latest_docx).paragraphs
    )
    assert "商业航天新闻周报" in "\n".join(
        paragraph.text for paragraph in Document(weekly_outputs.latest_docx).paragraphs
    )
