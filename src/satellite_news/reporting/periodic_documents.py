"""Generate the previous-day daily DOCX and the rolling seven-day weekly DOCX."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPORT_TIMEZONE = timezone(timedelta(hours=8), name="Asia/Shanghai")
DAILY_SCHEMA_VERSION = "daily_document.v1"
WEEKLY_SCHEMA_VERSION = "weekly_report.v1"
DEFAULT_CATALOG_PATH = Path("docs/data/news/archive/catalog.json")
DEFAULT_EVENT_TIMELINE_PATH = Path("docs/data/news/latest/event_timeline.json")
DEFAULT_DAILY_ROOT = Path("data/reports/daily-documents")
DEFAULT_DAILY_PUBLISH_ROOT = Path("docs/data/reports/daily-documents")
DEFAULT_WEEKLY_ROOT = Path("data/reports/weekly")
DEFAULT_WEEKLY_PUBLISH_ROOT = Path("docs/data/reports/weekly")
EAST_ASIA_FONT = "Kaiti SC"
EVENT_LABELS = {
    "launch": "发射与任务",
    "funding": "融资与资本",
    "order": "订单与合作",
    "regulation": "监管与政策",
    "market": "市场与股价",
    "product": "产品与技术",
    "other": "其他动态",
}
EVENT_ORDER = ("launch", "order", "funding", "regulation", "market", "product", "other")
EVENT_TYPE_ALIASES = {
    "financing": "funding",
    "partnership": "order",
    "corporate": "other",
}


@dataclass(frozen=True)
class DocumentOutputs:
    latest_docx: Path
    latest_json: Path
    published_docx: Path
    published_json: Path
    archived_docx: Path
    archived_json: Path


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def display_text(value: Any) -> str:
    text = re.sub(r"https?://\S+", "", str(value or ""))
    return re.sub(r"\s+", " ", text).strip(" -｜")


def parse_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(REPORT_TIMEZONE)


def local_date(value: str | None) -> date | None:
    parsed = parse_datetime(value)
    return parsed.date() if parsed else None


def resolve_as_of(payload: dict[str, Any], as_of: date | None) -> date:
    if as_of:
        return as_of
    generated = local_date(payload.get("generated_at"))
    return generated or datetime.now(REPORT_TIMEZONE).date()


def normalize_event_type(value: str | None) -> str:
    normalized = EVENT_TYPE_ALIASES.get(value or "", value or "other")
    return normalized if normalized in EVENT_LABELS else "other"


def item_event_type(item: dict[str, Any]) -> str:
    quality = item.get("quality") or {}
    metadata = item.get("metadata") or {}
    value = quality.get("event_type") or metadata.get("event_type")
    return normalize_event_type(value)


def item_source_name(item: dict[str, Any]) -> str:
    source = item.get("source") or {}
    return source.get("source_name") or source.get("source_id") or "未知来源"


def item_score(item: dict[str, Any]) -> float:
    quality = item.get("quality") or {}
    return float(quality.get("company_relevance_score") or 0) + float(
        quality.get("source_quality_score") or 0
    )


def item_sort_key(item: dict[str, Any]) -> tuple[float, str, str]:
    return (item_score(item), item.get("published_at") or "", item.get("id") or "")


def select_balanced_items(
    items: Iterable[dict[str, Any]],
    *,
    limit: int,
    per_company: int,
) -> list[dict[str, Any]]:
    selected: list[dict[str, Any]] = []
    company_counts: Counter[str] = Counter()
    for item in sorted(items, key=item_sort_key, reverse=True):
        company_id = item.get("company_id") or "unknown"
        if company_counts[company_id] >= per_company:
            continue
        selected.append(item)
        company_counts[company_id] += 1
        if len(selected) >= limit:
            break
    return selected


def compact_item(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "item_id": item.get("id"),
        "company_id": item.get("company_id"),
        "company_name": item.get("company_name") or item.get("company_id") or "未知公司",
        "title": item.get("title") or "无标题",
        "published_at": item.get("published_at"),
        "source_name": item_source_name(item),
        "event_type": item_event_type(item),
        "url": item.get("url"),
    }


def company_summary(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for item in items:
        grouped[item.get("company_id") or "unknown"].append(item)
    rows = []
    for company_id, company_items in grouped.items():
        company_items.sort(key=item_sort_key, reverse=True)
        rows.append(
            {
                "company_id": company_id,
                "company_name": company_items[0].get("company_name") or company_id,
                "item_count": len(company_items),
                "event_types": dict(Counter(item_event_type(item) for item in company_items)),
                "latest_title": company_items[0].get("title") or "无标题",
            }
        )
    return sorted(rows, key=lambda row: (-row["item_count"], row["company_name"]))


def build_previous_day_payload(
    catalog: dict[str, Any],
    *,
    as_of: date | None = None,
) -> dict[str, Any]:
    anchor = resolve_as_of(catalog, as_of)
    report_date = anchor - timedelta(days=1)
    all_items = catalog.get("items") or []
    day_items = [item for item in all_items if local_date(item.get("published_at")) == report_date]
    selected = select_balanced_items(day_items, limit=36, per_company=5)
    companies = company_summary(day_items)
    event_counts = Counter(item_event_type(item) for item in day_items)
    sources = Counter(item_source_name(item) for item in day_items)
    leading = "、".join(row["company_name"] for row in companies[:4])
    summary = (
        f"{report_date.isoformat()} 共归档 {len(day_items)} 条新闻，覆盖 {len(companies)} 家公司、"
        f"{len(sources)} 个信息来源。"
    )
    if leading:
        summary += f" 新闻量较多的公司包括 {leading}。"
    if not day_items:
        summary = f"{report_date.isoformat()} 未归档到新增新闻，历史新闻仍保留在新闻档案中。"
    return {
        "schema_version": DAILY_SCHEMA_VERSION,
        "document_type": "previous_day_daily_report",
        "generated_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "as_of_date": anchor.isoformat(),
        "report_date": report_date.isoformat(),
        "period_start": report_date.isoformat(),
        "period_end": report_date.isoformat(),
        "total_items": len(day_items),
        "company_count": len(companies),
        "source_count": len(sources),
        "summary": summary,
        "event_type_counts": dict(event_counts),
        "companies": companies,
        "top_items": [compact_item(item) for item in selected],
        "source_counts": dict(sources.most_common()),
    }


def event_articles_in_window(
    event: dict[str, Any],
    start: date,
    end: date,
) -> list[dict[str, Any]]:
    return [
        article
        for article in event.get("articles") or []
        if (published := local_date(article.get("published_at"))) and start <= published <= end
    ]


def event_sort_key(event: dict[str, Any]) -> tuple[float, str, int]:
    return (
        float(event.get("importance_score") or 0),
        event.get("latest_at") or "",
        int(event.get("source_count") or 0),
    )


def title_tokens(value: str) -> set[str]:
    generic_english = {
        "company",
        "launch",
        "launches",
        "mission",
        "news",
        "earth",
        "live",
        "orbit",
        "rocket",
        "satellite",
        "satellites",
        "space",
        "video",
        "watch",
        "the",
        "and",
        "with",
        "from",
    }
    generic_chinese = {
        "公司",
        "商业",
        "民营",
        "卫星",
        "航天",
        "火箭",
        "发射",
        "最新",
        "消息",
        "中国",
    }
    english = {
        token.lower()
        for token in re.findall(r"[A-Za-z0-9][A-Za-z0-9-]{2,}", value)
        if token.lower() not in generic_english
    }
    chinese_groups = re.findall(r"[\u4e00-\u9fff]{2,}", value)
    chinese_bigrams = {
        group[index : index + 2]
        for group in chinese_groups
        for index in range(max(0, len(group) - 1))
        if group[index : index + 2] not in generic_chinese
    }
    return english | chinese_bigrams


def relation_score(current: dict[str, Any], older: dict[str, Any]) -> float:
    if current.get("company_id") != older.get("company_id"):
        return 0
    if normalize_event_type(current.get("event_type")) != normalize_event_type(
        older.get("event_type")
    ):
        return 0
    current_tokens = title_tokens(current.get("headline") or "")
    older_tokens = title_tokens(older.get("headline") or "")
    company_tokens = title_tokens(current.get("company_name") or "") | title_tokens(
        str(current.get("company_id") or "").replace("_", " ")
    )
    current_tokens -= company_tokens
    older_tokens -= company_tokens
    if not current_tokens or not older_tokens:
        return 0
    intersection = current_tokens & older_tokens
    if len(intersection) < 3:
        return 0
    return len(intersection) / len(current_tokens | older_tokens)


def related_history(
    current_events: list[dict[str, Any]],
    all_events: list[dict[str, Any]],
    *,
    before: date,
) -> list[dict[str, Any]]:
    older_events = [
        event
        for event in all_events
        if (latest := local_date(event.get("latest_at"))) and latest < before
    ]
    contexts = []
    seen_pairs: set[tuple[str, str]] = set()
    for current in current_events:
        candidates = [
            (relation_score(current, older), older)
            for older in older_events
            if current.get("company_id") == older.get("company_id")
            and normalize_event_type(current.get("event_type"))
            == normalize_event_type(older.get("event_type"))
        ]
        candidates.sort(key=lambda pair: (pair[0], event_sort_key(pair[1])), reverse=True)
        if not candidates or candidates[0][0] < 0.2:
            continue
        score, older = candidates[0]
        pair_key = (current.get("event_id") or "", older.get("event_id") or "")
        if pair_key in seen_pairs:
            continue
        seen_pairs.add(pair_key)
        contexts.append(
            {
                "company_id": current.get("company_id"),
                "company_name": current.get("company_name") or current.get("company_id"),
                "event_type": normalize_event_type(current.get("event_type")),
                "current_event_id": current.get("event_id"),
                "current_headline": current.get("headline") or "无标题",
                "current_date": local_date(current.get("latest_at")).isoformat()
                if local_date(current.get("latest_at"))
                else None,
                "historical_event_id": older.get("event_id"),
                "historical_headline": older.get("headline") or "无标题",
                "historical_date": local_date(older.get("latest_at")).isoformat()
                if local_date(older.get("latest_at"))
                else None,
                "relation_basis": "同公司、同事件类别且标题主题相近",
                "relation_score": round(score, 3),
            }
        )
        if len(contexts) >= 6:
            break
    return contexts


def compact_event(
    event: dict[str, Any],
    *,
    weekly_articles: list[dict[str, Any]],
) -> dict[str, Any]:
    return {
        "event_id": event.get("event_id"),
        "company_id": event.get("company_id"),
        "company_name": event.get("company_name") or event.get("company_id") or "未知公司",
        "event_type": normalize_event_type(event.get("event_type")),
        "event_label": event.get("event_label")
        or EVENT_LABELS.get(normalize_event_type(event.get("event_type")), "其他动态"),
        "headline": event.get("headline") or "无标题",
        "summary": event.get("summary") or "",
        "started_at": event.get("started_at"),
        "latest_at": event.get("latest_at"),
        "importance_score": event.get("importance_score") or 0,
        "source_count": event.get("source_count") or 0,
        "weekly_article_count": len(weekly_articles),
        "source_names": event.get("source_names") or [],
        "latest_url": event.get("latest_url"),
    }


def build_weekly_payload(
    catalog: dict[str, Any],
    event_timeline: dict[str, Any],
    *,
    as_of: date | None = None,
) -> dict[str, Any]:
    anchor = resolve_as_of(catalog, as_of)
    period_start = anchor - timedelta(days=6)
    all_items = catalog.get("items") or []
    week_items = [
        item
        for item in all_items
        if (published := local_date(item.get("published_at"))) and period_start <= published <= anchor
    ]
    companies = company_summary(week_items)
    sources = Counter(item_source_name(item) for item in week_items)
    all_events = event_timeline.get("events") or []
    event_pairs = [
        (event, weekly_articles)
        for event in all_events
        if (weekly_articles := event_articles_in_window(event, period_start, anchor))
    ]
    event_pairs.sort(key=lambda pair: event_sort_key(pair[0]), reverse=True)
    selected_pairs = event_pairs[:30]
    selected_events = [pair[0] for pair in selected_pairs]
    events = [
        compact_event(event, weekly_articles=weekly_articles)
        for event, weekly_articles in selected_pairs
    ]
    contexts = related_history(selected_events, all_events, before=period_start)
    event_counts = Counter(
        normalize_event_type(event.get("event_type")) for event, _ in event_pairs
    )
    leading = "、".join(row["company_name"] for row in companies[:5])
    summary = (
        f"{period_start.isoformat()} 至 {anchor.isoformat()} 共归档 {len(week_items)} 条新闻，"
        f"形成 {len(event_pairs)} 个事件，覆盖 {len(companies)} 家公司和 {len(sources)} 个来源。"
    )
    if leading:
        summary += f" 新闻量较多的公司包括 {leading}。"
    if contexts:
        summary += f" 其中 {len(contexts)} 项动态可与更早的同公司同类事件关联回看。"
    if not week_items:
        summary = f"{period_start.isoformat()} 至 {anchor.isoformat()} 未归档到新增新闻。"
    return {
        "schema_version": WEEKLY_SCHEMA_VERSION,
        "document_type": "weekly_news_report",
        "generated_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "as_of_date": anchor.isoformat(),
        "period_start": period_start.isoformat(),
        "period_end": anchor.isoformat(),
        "total_items": len(week_items),
        "event_count": len(event_pairs),
        "company_count": len(companies),
        "source_count": len(sources),
        "summary": summary,
        "event_type_counts": dict(event_counts),
        "companies": companies,
        "events": events,
        "related_history": contexts,
        "source_counts": dict(sources.most_common()),
    }


def set_run_font(run: Any, *, size: float = 10.5, bold: bool | None = None) -> None:
    run.font.name = EAST_ASIA_FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    fonts = run._element.get_or_add_rPr().rFonts
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), EAST_ASIA_FONT)


def set_cell_text(cell: Any, text: str, *, bold: bool = False, align: str = "left") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(display_text(text))
    set_run_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=8)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def configure_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal = document.styles["Normal"]
    normal.font.name = EAST_ASIA_FONT
    normal.font.size = Pt(10.5)
    normal_fonts = normal._element.get_or_add_rPr().rFonts
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_fonts.set(qn(f"w:{attribute}"), EAST_ASIA_FONT)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in (("Title", 18), ("Heading 1", 12), ("Heading 2", 10.5)):
        style = document.styles[style_name]
        style.font.name = EAST_ASIA_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style_fonts = style._element.get_or_add_rPr().rFonts
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            style_fonts.set(qn(f"w:{attribute}"), EAST_ASIA_FONT)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
    document.styles["Heading 1"].paragraph_format.page_break_before = False

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return document


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(title)
    set_run_font(run, size=18, bold=True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=10.5)


def add_heading(document: Document, number: int, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(f"{number}. {title}")
    set_run_font(run, size=10.5, bold=True)


def add_bullet(document: Document, text: str, *, marker: str = "➢") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.first_line_indent = Cm(-0.4)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_after = Pt(3)
    marker_run = paragraph.add_run(f"{marker} ")
    set_run_font(marker_run, size=10.5)
    run = paragraph.add_run(display_text(text))
    set_run_font(run, size=10.5)


def add_metadata_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row, (label, value) in zip(table.rows, rows):
        prevent_row_split(row)
        set_cell_text(row.cells[0], label, bold=True, align="center")
        set_cell_text(row.cells[1], value)
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(11.4)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_company_table(document: Document, companies: list[dict[str, Any]], *, limit: int = 15) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("公司", "新闻数", "事件类别", "最新动态")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, align="center")
    set_repeat_table_header(table.rows[0])
    for company in companies[:limit]:
        row = table.add_row()
        prevent_row_split(row)
        labels = "、".join(
            EVENT_LABELS.get(event_type, event_type)
            for event_type in company.get("event_types") or {}
        )
        values = (
            company.get("company_name") or "未知公司",
            str(company.get("item_count") or 0),
            labels or "其他动态",
            company.get("latest_title") or "无标题",
        )
        for index, (cell, text) in enumerate(zip(row.cells, values)):
            set_cell_text(cell, text, align="center" if index in (1, 2) else "left")


def add_source_appendix(
    document: Document,
    items: list[dict[str, Any]],
    *,
    limit: int = 30,
) -> None:
    for index, item in enumerate(items[:limit], start=1):
        published = local_date(item.get("published_at"))
        date_text = published.isoformat() if published else "日期未知"
        add_bullet(
            document,
            f"[{index}] {date_text}｜{item.get('source_name') or '未知来源'}｜"
            f"{item.get('company_name') or '未知公司'}：{item.get('title') or '无标题'}",
            marker="✓",
        )


def render_daily_document(payload: dict[str, Any], path: Path) -> None:
    document = configure_document()
    add_title(document, "商业航天新闻日报", f"{payload['report_date']}（上一自然日）")
    add_metadata_table(
        document,
        [
            ("统计日期", payload["report_date"]),
            ("新闻 / 公司", f"{payload['total_items']} 条 / {payload['company_count']} 家"),
            ("信息来源", f"{payload['source_count']} 个"),
            ("生成日期", payload["as_of_date"]),
        ],
    )
    add_heading(document, 1, "昨日摘要")
    document.add_paragraph(payload["summary"])

    add_heading(document, 2, "重点动态")
    if payload["top_items"]:
        for item in payload["top_items"][:18]:
            event_label = EVENT_LABELS.get(item.get("event_type") or "other", "其他动态")
            add_bullet(
                document,
                f"【{event_label}】{item['company_name']}：{item['title']} "
                f"（{item['source_name']}）",
            )
    else:
        add_bullet(document, "昨日未归档到新增新闻。")

    add_heading(document, 3, "公司动态概览")
    if payload["companies"]:
        add_company_table(document, payload["companies"], limit=15)
    else:
        document.add_paragraph("无公司动态。")

    add_heading(document, 4, "来源附录")
    add_source_appendix(document, payload["top_items"], limit=30)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def add_event_table(document: Document, events: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("日期", "公司", "类别", "核心事件")):
        set_cell_text(cell, text, bold=True, align="center")
    set_repeat_table_header(table.rows[0])
    for event in events:
        row = table.add_row()
        prevent_row_split(row)
        published = local_date(event.get("latest_at"))
        values = (
            published.isoformat() if published else "日期未知",
            event.get("company_name") or "未知公司",
            EVENT_LABELS.get(event.get("event_type") or "other", "其他动态"),
            event.get("headline") or "无标题",
        )
        for index, (cell, text) in enumerate(zip(row.cells, values)):
            set_cell_text(cell, text, align="center" if index < 3 else "left")


def render_weekly_document(payload: dict[str, Any], path: Path) -> None:
    document = configure_document()
    subtitle = f"{payload['period_start']} 至 {payload['period_end']}"
    add_title(document, "商业航天新闻周报", subtitle)
    add_metadata_table(
        document,
        [
            ("统计周期", subtitle),
            ("新闻 / 事件", f"{payload['total_items']} 条 / {payload['event_count']} 个"),
            ("公司 / 来源", f"{payload['company_count']} 家 / {payload['source_count']} 个"),
            ("生成日期", payload["as_of_date"]),
        ],
    )

    add_heading(document, 1, "本周摘要")
    document.add_paragraph(payload["summary"])
    for event_type in EVENT_ORDER:
        count = payload["event_type_counts"].get(event_type, 0)
        if count:
            add_bullet(document, f"{EVENT_LABELS[event_type]}：{count} 个事件。", marker="✓")

    add_heading(document, 2, "核心事件")
    if payload["events"]:
        add_event_table(document, payload["events"][:18])
    else:
        document.add_paragraph("本周未形成可展示事件。")

    section_number = 3
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for event in payload["events"]:
        grouped[event.get("event_type") or "other"].append(event)
    for event_type in EVENT_ORDER:
        events = grouped.get(event_type) or []
        if not events:
            continue
        add_heading(document, section_number, EVENT_LABELS[event_type])
        section_number += 1
        for event in events[:8]:
            latest = local_date(event.get("latest_at"))
            source_names = "、".join((event.get("source_names") or [])[:3]) or "来源未标注"
            detail = (
                f"{event['company_name']}：{event['headline']} "
                f"（{latest.isoformat() if latest else '日期未知'}；{source_names}）"
            )
            add_bullet(document, detail)

    add_heading(document, section_number, "公司动态概览")
    section_number += 1
    if payload["companies"]:
        add_company_table(document, payload["companies"], limit=18)
    else:
        document.add_paragraph("本周无公司动态。")

    add_heading(document, section_number, "关联历史回看")
    section_number += 1
    if payload["related_history"]:
        for context in payload["related_history"]:
            add_bullet(
                document,
                f"{context['company_name']}：本周“{context['current_headline']}”；"
                f"关联回看 {context['historical_date']} 的“{context['historical_headline']}”。",
                marker="✓",
            )
        document.add_paragraph(
            "说明：以上按同公司、同事件类别及标题主题相近建立线索，仅供连续观察，"
            "不代表已确认因果关系。"
        )
    else:
        document.add_paragraph("本周未识别到主题相近且可可靠回看的历史事件。")

    add_heading(document, section_number, "下周关注")
    if payload["events"]:
        leading_types = [
            EVENT_LABELS.get(event_type, event_type)
            for event_type, _ in Counter(
                event.get("event_type") or "other" for event in payload["events"]
            ).most_common(3)
        ]
        add_bullet(document, f"继续跟踪本周高频类别：{'、'.join(leading_types)}。")
        for company in payload["companies"][:5]:
            add_bullet(
                document,
                f"关注 {company['company_name']} 的后续披露与多来源交叉验证。",
                marker="✓",
            )
    else:
        add_bullet(document, "保持对新增发射、融资、订单、监管及市场信息的跟踪。")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def write_document_outputs(
    *,
    payload: dict[str, Any],
    local_root: Path,
    publish_root: Path,
    report_date: date,
    stable_name: str,
    archive_name: str,
    renderer: Any,
) -> DocumentOutputs:
    latest_docx = local_root / "latest" / stable_name
    latest_json = latest_docx.with_suffix(".json")
    published_docx = publish_root / "latest" / stable_name
    published_json = published_docx.with_suffix(".json")
    archive_dir = local_root / "archive" / report_date.strftime("%Y/%m/%d")
    archived_docx = archive_dir / archive_name
    archived_json = archived_docx.with_suffix(".json")

    renderer(payload, latest_docx)
    write_json(latest_json, payload)
    published_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(latest_docx, published_docx)
    write_json(published_json, payload)
    archive_dir.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(latest_docx, archived_docx)
    write_json(archived_json, payload)
    return DocumentOutputs(
        latest_docx,
        latest_json,
        published_docx,
        published_json,
        archived_docx,
        archived_json,
    )


def generate_previous_day_document(
    *,
    catalog_path: Path = DEFAULT_CATALOG_PATH,
    local_root: Path = DEFAULT_DAILY_ROOT,
    publish_root: Path = DEFAULT_DAILY_PUBLISH_ROOT,
    as_of: date | None = None,
) -> tuple[dict[str, Any], DocumentOutputs]:
    payload = build_previous_day_payload(load_json(catalog_path), as_of=as_of)
    report_date = date.fromisoformat(payload["report_date"])
    outputs = write_document_outputs(
        payload=payload,
        local_root=local_root,
        publish_root=publish_root,
        report_date=report_date,
        stable_name="previous_daily_report.docx",
        archive_name=f"商业航天新闻日报_{report_date.isoformat()}.docx",
        renderer=render_daily_document,
    )
    return payload, outputs


def generate_weekly_document(
    *,
    catalog_path: Path = DEFAULT_CATALOG_PATH,
    event_timeline_path: Path = DEFAULT_EVENT_TIMELINE_PATH,
    local_root: Path = DEFAULT_WEEKLY_ROOT,
    publish_root: Path = DEFAULT_WEEKLY_PUBLISH_ROOT,
    as_of: date | None = None,
) -> tuple[dict[str, Any], DocumentOutputs]:
    catalog = load_json(catalog_path)
    timeline = load_json(event_timeline_path)
    payload = build_weekly_payload(catalog, timeline, as_of=as_of)
    period_end = date.fromisoformat(payload["period_end"])
    outputs = write_document_outputs(
        payload=payload,
        local_root=local_root,
        publish_root=publish_root,
        report_date=period_end,
        stable_name="weekly_report.docx",
        archive_name=(
            f"商业航天新闻周报_{payload['period_start']}_{payload['period_end']}.docx"
        ),
        renderer=render_weekly_document,
    )
    return payload, outputs


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--kind", choices=("daily", "weekly", "all"), default="all")
    parser.add_argument("--catalog", type=Path, default=DEFAULT_CATALOG_PATH)
    parser.add_argument("--event-timeline", type=Path, default=DEFAULT_EVENT_TIMELINE_PATH)
    parser.add_argument("--as-of", type=date.fromisoformat)
    parser.add_argument("--daily-root", type=Path, default=DEFAULT_DAILY_ROOT)
    parser.add_argument("--daily-publish-root", type=Path, default=DEFAULT_DAILY_PUBLISH_ROOT)
    parser.add_argument("--weekly-root", type=Path, default=DEFAULT_WEEKLY_ROOT)
    parser.add_argument("--weekly-publish-root", type=Path, default=DEFAULT_WEEKLY_PUBLISH_ROOT)
    args = parser.parse_args(argv)

    if args.kind in ("daily", "all"):
        payload, outputs = generate_previous_day_document(
            catalog_path=args.catalog,
            local_root=args.daily_root,
            publish_root=args.daily_publish_root,
            as_of=args.as_of,
        )
        print(
            f"Generated previous-day report date={payload['report_date']} "
            f"items={payload['total_items']} docx={outputs.latest_docx}"
        )
    if args.kind in ("weekly", "all"):
        payload, outputs = generate_weekly_document(
            catalog_path=args.catalog,
            event_timeline_path=args.event_timeline,
            local_root=args.weekly_root,
            publish_root=args.weekly_publish_root,
            as_of=args.as_of,
        )
        print(
            f"Generated weekly report period={payload['period_start']}..{payload['period_end']} "
            f"items={payload['total_items']} docx={outputs.latest_docx}"
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
